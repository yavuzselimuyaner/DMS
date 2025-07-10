from flask import Flask, request, redirect, url_for, render_template, flash, send_file, session
import os
from werkzeug.utils import secure_filename
import json
from flask_sqlalchemy import SQLAlchemy
from io import BytesIO
import datetime
from werkzeug.security import generate_password_hash, check_password_hash
import base64
import secrets
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
METADATA_FILE = 'uploads/metadata.json'

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'your_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:yavuz@localhost:3306/dms2'

# E-mail configuration
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'yavuzselimuyaner@gmail.com'  # Sizin e-posta adresiniz
app.config['MAIL_PASSWORD'] = 'iaamgssfqdrkftdz'          # Gmail app password
app.config['MAIL_DEFAULT_SENDER'] = 'yavuzselimuyaner@gmail.com'

db = SQLAlchemy(app)
# Simple user system (for demo)
USERS = {
    'admin': {'password': 'adminpass', 'role': 'admin'},
    'user': {'password': 'userpass', 'role': 'user'}
}

def is_admin():
    return session.get('role') == 'admin'

def is_logged_in():
    return 'username' in session

def get_accessible_document_types():
    """Get document types that current user can access based on hierarchy"""
    user_role = session.get('role', 'user')
    
    # Hierarchy: admin > user > calisan (çalışan)
    if user_role == 'admin':
        # Admin can see all document types
        return DocumentType.query.all()
    elif user_role == 'user':
        # User can see user and employee documents
        return DocumentType.query.filter(DocumentType.name.in_(['Kullanıcı', 'Çalışan'])).all()
    elif user_role == 'calisan':
        # Employee can only see employee documents
        return DocumentType.query.filter_by(name='Çalışan').all()
    else:
        # Default: can only see employee documents
        return DocumentType.query.filter_by(name='Çalışan').all()

def can_access_document(document):
    """Check if current user can access this document based on hierarchy"""
    user_role = session.get('role', 'user')
    
    if not document.document_type_id:
        return True  # Old documents without type are accessible to all
    
    doc_type = DocumentType.query.get(document.document_type_id)
    if not doc_type:
        return True
    
    # Hierarchy check
    if user_role == 'admin':
        return True  # Admin can access all
    elif user_role == 'user':
        return doc_type.name in ['Kullanıcı', 'Çalışan']
    elif user_role == 'calisan':
        return doc_type.name == 'Çalışan'
    else:
        # Default: can only see employee documents
        return doc_type.name == 'Çalışan'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_thumbnail(file_data, file_ext):
    """Generate thumbnail for different file types"""
    try:
        if file_ext == '.pdf':
            # PDF thumbnail using PyMuPDF (fitz)
            try:
                import fitz  # PyMuPDF
                pdf_document = fitz.open(stream=file_data, filetype="pdf")
                first_page = pdf_document[0]
                pix = first_page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))  # Scale down
                img_data = pix.tobytes("png")
                pdf_document.close()
                return base64.b64encode(img_data).decode('utf-8')
            except ImportError:
                return None
        elif file_ext == '.txt':
            # Text thumbnail - first few lines
            try:
                content = file_data.decode('utf-8', errors='ignore')
                lines = content.splitlines()[:8]  # First 8 lines
                preview_text = '\n'.join(lines)
                return preview_text[:300]  # First 300 chars
            except:
                return None
        elif file_ext in ['.docx', '.doc']:
            # Word document thumbnail - extract first paragraph
            try:
                from docx import Document as DocxDocument
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                    tmp.write(file_data)
                    tmp_path = tmp.name
                docx_doc = DocxDocument(tmp_path)
                paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
                preview_text = '\n'.join(paragraphs[:3])  # First 3 paragraphs
                os.remove(tmp_path)
                return preview_text[:300]  # First 300 chars
            except:
                return None
    except Exception:
        return None
    return None

def load_metadata():
    if os.path.exists(METADATA_FILE):
        with open(METADATA_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_metadata(metadata):
    with open(METADATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if 'login' in request.form:
            # Giriş işlemi
            username = request.form['username']
            password = request.form['password']
            user = db.session.query(User).filter_by(username=username).first()
            if user and check_password_hash(user.password_hash, password):
                session['username'] = user.username
                session['role'] = user.role
                flash('Giriş başarılı!')
                return redirect(url_for('upload_file'))
            else:
                flash('Geçersiz kullanıcı adı veya şifre')
        elif 'register' in request.form:
            # Kayıt işlemi
            username = request.form['reg_username'].strip()
            email = request.form['reg_email'].strip()
            password = request.form['reg_password']
            # Kullanıcı adı veya e-posta var mı?
            if db.session.query(User).filter((User.username == username) | (User.email == email)).first():
                flash('Kullanıcı adı veya e-posta zaten mevcut')
            else:
                user = User(
                    username=username,
                    email=email,
                    password_hash=generate_password_hash(password),
                    role='user',
                    created_at=datetime.datetime.now()
                )
                db.session.add(user)
                db.session.commit()
                flash('Hesap oluşturuldu! Şimdi giriş yapabilirsiniz.')
    return render_template('login_register.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Çıkış yapıldı.')
    return redirect(url_for('login'))

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if not is_logged_in():
        return redirect(url_for('login'))
    
    user = db.session.query(User).filter_by(username=session['username']).first()
    if not user:
        flash('Kullanıcı bulunamadı')
        return redirect(url_for('login'))
    
    # Get user statistics
    user_documents = Document.query.filter_by(uploaded_by=user.id).all()
    total_documents = len(user_documents)
    
    # Calculate total storage used (approximate)
    total_size_bytes = sum(len(doc.file_data) if doc.file_data else 0 for doc in user_documents)
    total_size = f"{total_size_bytes / (1024*1024):.1f} MB" if total_size_bytes > 0 else "0 MB"
    
    # Last upload date
    last_upload = "Hiç yüklenmedi"
    if user_documents:
        latest_doc = max(user_documents, key=lambda x: x.upload_date if x.upload_date else datetime.datetime.min)
        if latest_doc.upload_date:
            last_upload = latest_doc.upload_date.strftime('%d %B %Y')
    
    user_stats = {
        'total_documents': total_documents,
        'total_size': total_size,
        'last_upload': last_upload
    }
    
    # Admin-specific statistics
    admin_stats = None
    all_documents = []
    all_users = []
    if is_admin():
        # Get all system statistics for admin
        all_documents = Document.query.all()
        all_users = User.query.all()
        total_system_docs = len(all_documents)
        total_system_size_bytes = sum(len(doc.file_data) if doc.file_data else 0 for doc in all_documents)
        total_system_size = f"{total_system_size_bytes / (1024*1024):.1f} MB" if total_system_size_bytes > 0 else "0 MB"
        total_users = len(all_users)
        
        admin_stats = {
            'total_system_documents': total_system_docs,
            'total_system_size': total_system_size,
            'total_users': total_users,
            'recent_uploads': sorted(all_documents, key=lambda x: x.upload_date if x.upload_date else datetime.datetime.min, reverse=True)[:5]
        }
        
        # Create user map for admin view
        user_map = {u.id: u.username for u in all_users}
        user_map[None] = 'Sistem Kullanıcısı'
        
        # Add user info to documents for admin view
        for doc in all_documents:
            doc.uploader_name = user_map.get(doc.uploaded_by, 'Bilinmeyen Kullanıcı')
    
    if request.method == 'POST':
        current_password = request.form['current_password']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']
        
        # Verify current password
        if not check_password_hash(user.password_hash, current_password):
            flash('Mevcut şifre yanlış')
            return render_template('profile.html', user=user, user_stats=user_stats, admin_stats=admin_stats, all_documents=all_documents, all_users=all_users, is_admin=is_admin())
        
        # Check if new passwords match
        if new_password != confirm_password:
            flash('Yeni şifreler eşleşmiyor')
            return render_template('profile.html', user=user, user_stats=user_stats, admin_stats=admin_stats, all_documents=all_documents, all_users=all_users, is_admin=is_admin())
        
        # Check password length
        if len(new_password) < 6:
            flash('Şifre en az 6 karakter olmalıdır')
            return render_template('profile.html', user=user, user_stats=user_stats, admin_stats=admin_stats, all_documents=all_documents, all_users=all_users, is_admin=is_admin())
        
        # Update password
        user.password_hash = generate_password_hash(new_password)
        db.session.commit()
        flash('Şifre başarıyla güncellendi!')
        return render_template('profile.html', user=user, user_stats=user_stats, admin_stats=admin_stats, all_documents=all_documents, all_users=all_users, is_admin=is_admin())
    
    return render_template('profile.html', user=user, user_stats=user_stats, admin_stats=admin_stats, all_documents=all_documents, all_users=all_users, is_admin=is_admin())

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if not is_logged_in():
        return redirect(url_for('login'))
    
    # Get view mode - either upload form or document list
    view_mode = request.args.get('view', 'upload')  # 'upload' or 'documents'
    
    query = request.args.get('q', '').lower()
    sort_by = request.args.get('sort', 'name')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    page = request.args.get('page', 1, type=int)
    per_page = 9
    
    # Get documents with hierarchy filtering
    documents = Document.query
    
    # Apply hierarchy filter
    accessible_doc_types = get_accessible_document_types()
    accessible_type_ids = [dt.id for dt in accessible_doc_types]
    
    if accessible_type_ids:
        # Include documents with accessible types OR documents without type (old documents)
        documents = documents.filter(
            db.or_(
                Document.document_type_id.in_(accessible_type_ids),
                Document.document_type_id.is_(None)
            )
        )
    else:
        # If no accessible types, only show documents without type
        documents = documents.filter(Document.document_type_id.is_(None))
    
    if query:
        documents = documents.filter(Document.title.ilike(f'%{query}%'))
    if date_from:
        from_dt = datetime.datetime.strptime(date_from, '%Y-%m-%d')
        documents = documents.filter(Document.upload_date >= from_dt)
    if date_to:
        to_dt = datetime.datetime.strptime(date_to, '%Y-%m-%d')
        to_dt = to_dt.replace(hour=23, minute=59, second=59)
        documents = documents.filter(Document.upload_date <= to_dt)
    if sort_by == 'date':
        documents = documents.order_by(Document.upload_date.desc())
    else:
        documents = documents.order_by(Document.title.asc())
    documents = documents.all()
    
    # Pagination
    total = len(documents)
    page_count = (total // per_page) + (1 if total % per_page else 0)
    start = (page - 1) * per_page
    end = start + per_page
    paginated_documents = documents[start:end]
    
    # Get document types for upload form
    document_types = DocumentType.query.all()
    
    if request.method == 'POST':
        files = request.files.getlist('file')
        explanation = request.form.get('explanation', '').strip()
        document_type_id = request.form.get('document_type_id')
        
        if not files or files[0].filename == '':
            flash('Dosya seçilmedi')
            return redirect(request.url)
        
        # Convert document_type_id to int if provided
        if document_type_id and document_type_id != '':
            try:
                document_type_id = int(document_type_id)
            except ValueError:
                document_type_id = None
        else:
            document_type_id = None
        
        user = None
        if 'username' in session:
            user = db.session.query(User).filter_by(username=session['username']).first()
        uploaded_count = 0
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_data = file.read()
                doc = Document(
                    title=filename,
                    description=explanation,
                    upload_date=datetime.datetime.now(),
                    file_path='',
                    uploaded_by=user.id if user else None,
                    document_type_id=document_type_id,
                    access_level='private',
                    file_data=file_data
                )
                db.session.add(doc)
                uploaded_count += 1
        if uploaded_count > 0:
            db.session.commit()
            flash(f'{uploaded_count} dosya başarıyla yüklendi')
        else:
            flash('Geçerli dosya yüklenmedi')
        return redirect(url_for('upload_file'))
    
    # Create user map and document type map
    user_map = {}
    user_ids = set(doc.uploaded_by for doc in documents if doc.uploaded_by)
    if user_ids:
        users = User.query.filter(User.id.in_(user_ids)).all()
        user_map = {u.id: u.username for u in users}
    user_map[None] = 'Sistem Kullanıcısı'
    
    # Create document type map
    doc_type_map = {}
    doc_type_ids = set(doc.document_type_id for doc in documents if doc.document_type_id)
    if doc_type_ids:
        doc_types = DocumentType.query.filter(DocumentType.id.in_(doc_type_ids)).all()
        doc_type_map = {dt.id: dt.name for dt in doc_types}
    doc_type_map[None] = 'Tür Belirtilmemiş'
    
    # Generate thumbnails for documents
    thumbnails = {}
    previews = {}
    for doc in documents:
        ext = os.path.splitext(doc.title)[1].lower()
        if doc.file_data:
            thumbnail = generate_thumbnail(doc.file_data, ext)
            thumbnails[doc.id] = thumbnail
        else:
            thumbnails[doc.id] = None
            
        if ext == '.txt' and doc.file_data:
            try:
                content = doc.file_data.decode('utf-8', errors='ignore')
                previews[doc.id] = '\n'.join(content.splitlines()[:3])
            except Exception:
                previews[doc.id] = ''
        else:
            previews[doc.id] = ''
    
    return render_template('index.html', 
                         documents=documents, 
                         paginated_documents=paginated_documents, 
                         sort_by=sort_by, 
                         query=query, 
                         date_from=date_from, 
                         date_to=date_to, 
                         is_admin=is_admin(), 
                         previews=previews, 
                         page=page, 
                         per_page=per_page, 
                         page_count=page_count, 
                         total=total, 
                         user_map=user_map, 
                         thumbnails=thumbnails,
                         document_types=document_types,
                         doc_type_map=doc_type_map,
                         view_mode=view_mode)

@app.route('/delete/<int:doc_id>', methods=['POST'])
def delete_file(doc_id):
    if not is_admin():
        flash('Yetkisiz erişim')
        return redirect(url_for('upload_file'))
    doc = Document.query.get(doc_id)
    if doc:
        db.session.delete(doc)
        db.session.commit()
        flash('Dosya silindi')
    else:
        flash('Dosya bulunamadı')
    return redirect(url_for('upload_file'))

@app.route('/download/<int:doc_id>')
def download_file(doc_id):
    doc = Document.query.get(doc_id)
    if not doc or not doc.file_data:
        flash('Dosya bulunamadı')
        return redirect(url_for('upload_file'))
    # PDF önizleme için as_attachment parametresi kontrolü
    preview = request.args.get('preview', '0') == '1'
    return send_file(BytesIO(doc.file_data), download_name=doc.title, as_attachment=not preview)

@app.route('/preview/<int:doc_id>')
def preview_file(doc_id):
    if not is_logged_in():
        return redirect(url_for('login'))
    doc = Document.query.get(doc_id)
    if not doc or not doc.file_data:
        flash('Dosya bulunamadı')
        return redirect(url_for('upload_file'))
    ext = os.path.splitext(doc.title)[1].lower()
    if ext == '.pdf':
        return render_template('preview_pdf.html', filename=doc.title, doc_id=doc.id)
    elif ext == '.txt':
        content = doc.file_data.decode('utf-8', errors='ignore')
        return render_template('preview_txt.html', filename=doc.title, content=content, doc_id=doc.id)
    elif ext == '.docx':
        try:
            from docx import Document as DocxDocument
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(doc.file_data)
                tmp_path = tmp.name
            docx_doc = DocxDocument(tmp_path)
            paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
            content = '\n'.join(paragraphs)
            os.remove(tmp_path)
        except Exception as e:
            content = f'DOCX okuma hatası: {e}'
        return render_template('preview_docx.html', filename=doc.title, content=content, doc_id=doc.id)
    else:
        flash('Bu dosya türü için önizleme desteklenmiyor.')
        return redirect(url_for('upload_file'))

@app.route('/register_only', methods=['GET', 'POST'])
def register_only():
    if request.method == 'POST':
        username = request.form['reg_username'].strip()
        email = request.form['reg_email'].strip()
        password = request.form['reg_password']
        password2 = request.form['reg_password2']
        role = request.form.get('reg_role', 'user').strip()
        
        # Rol validasyonu
        if role not in ['user', 'calisan']:
            role = 'user'
        
        if password != password2:
            flash('Şifreler eşleşmiyor!')
            return render_template('register_only.html')
        if db.session.query(User).filter((User.username == username) | (User.email == email)).first():
            flash('Kullanıcı adı veya e-posta zaten mevcut')
            return render_template('register_only.html')
        user = User(
            username=username,
            email=email,
            password_hash=generate_password_hash(password),
            role=role,
            created_at=datetime.datetime.now()
        )
        db.session.add(user)
        db.session.commit()
        flash('Hesap oluşturuldu! Şimdi giriş yapabilirsiniz.')
        return redirect(url_for('login'))
    return render_template('register_only.html')

@app.route('/init_admin')
def init_admin():
    """Initialize default admin user and document types if not exists"""
    # Create all database tables
    with app.app_context():
        db.create_all()
    
    # Create admin user
    admin_user = User.query.filter_by(username='admin').first()
    if not admin_user:
        admin_user = User(
            username='admin',
            email='admin@dms.local',
            password_hash=generate_password_hash('adminpass'),
            role='admin',
            created_at=datetime.datetime.now()
        )
        db.session.add(admin_user)
    
    # Create document types
    doc_types = ['Yönetici', 'Kullanıcı', 'Çalışan']
    created_types = []
    for doc_type_name in doc_types:
        existing_type = DocumentType.query.filter_by(name=doc_type_name).first()
        if not existing_type:
            doc_type = DocumentType(name=doc_type_name)
            db.session.add(doc_type)
            created_types.append(doc_type_name)
    
    try:
        db.session.commit()
        result_msg = "Kurulum tamamlandı!\n"
        if not User.query.filter_by(username='admin').first():
            result_msg += "✅ Admin kullanıcısı oluşturuldu (Kullanıcı adı: admin, Şifre: adminpass)\n"
        else:
            result_msg += "ℹ️ Admin kullanıcısı zaten mevcut\n"
        
        if created_types:
            result_msg += f"✅ Döküman türleri oluşturuldu: {', '.join(created_types)}\n"
        else:
            result_msg += "ℹ️ Döküman türleri zaten mevcut\n"
            
        result_msg += "✅ Veritabanı tabloları kontrol edildi ve oluşturuldu"
        return result_msg.replace('\n', '<br>')
    except Exception as e:
        db.session.rollback()
        return f"Hata oluştu: {str(e)}"

# SQLAlchemy Models
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.Text, nullable=False)
    role = db.Column(db.String(20), nullable=False)
    created_at = db.Column(db.DateTime)
    reset_token = db.Column(db.String(100), nullable=True)
    reset_token_expires = db.Column(db.DateTime, nullable=True)

class DocumentType(db.Model):
    __tablename__ = 'document_types'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)

class Document(db.Model):
    __tablename__ = 'documents'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text)
    upload_date = db.Column(db.DateTime)
    file_path = db.Column(db.Text, nullable=False)
    uploaded_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    document_type_id = db.Column(db.Integer, db.ForeignKey('document_types.id'))
    access_level = db.Column(db.String(20), default='private')
    file_data = db.Column(db.LargeBinary)  # <-- YENİ: Dosya verisi (BLOB)

class Permission(db.Model):
    __tablename__ = 'permissions'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id', ondelete='CASCADE'))
    document_id = db.Column(db.Integer, db.ForeignKey('documents.id', ondelete='CASCADE'))
    can_view = db.Column(db.Boolean, default=True)
    can_edit = db.Column(db.Boolean, default=False)
    can_delete = db.Column(db.Boolean, default=False)
    __table_args__ = (db.UniqueConstraint('user_id', 'document_id'),)

@app.route('/admin/delete_document/<int:doc_id>', methods=['POST'])
def admin_delete_document(doc_id):
    if not is_admin():
        flash('Yetkisiz erişim')
        return redirect(url_for('profile'))
    
    doc = Document.query.get(doc_id)
    if doc:
        db.session.delete(doc)
        db.session.commit()
        flash(f'Döküman "{doc.title}" başarıyla silindi')
    else:
        flash('Döküman bulunamadı')
    
    return redirect(url_for('profile'))

@app.route('/admin/delete_user/<int:user_id>', methods=['POST'])
def admin_delete_user(user_id):
    if not is_admin():
        flash('Yetkisiz erişim')
        return redirect(url_for('profile'))
    
    # Prevent admin from deleting themselves
    current_user = db.session.query(User).filter_by(username=session['username']).first()
    if current_user and current_user.id == user_id:
        flash('Kendi hesabınızı silemezsiniz')
        return redirect(url_for('profile'))
    
    user = User.query.get(user_id)
    if user:
        # Delete user's documents first
        user_docs = Document.query.filter_by(uploaded_by=user.id).all()
        for doc in user_docs:
            db.session.delete(doc)
        
        # Delete the user
        db.session.delete(user)
        db.session.commit()
        flash(f'Kullanıcı "{user.username}" ve dökümanları başarıyla silindi')
    else:
        flash('Kullanıcı bulunamadı')
    
    return redirect(url_for('profile'))

@app.route('/admin/promote_user/<int:user_id>', methods=['POST'])
def admin_promote_user(user_id):
    if not is_admin():
        flash('Unauthorized access')
        return redirect(url_for('profile'))
    
    user = User.query.get(user_id)
    if user:
        if user.role == 'user':
            user.role = 'admin'
            db.session.commit()
            flash(f'Kullanıcı "{user.username}" yönetici olarak yükseltildi')
        elif user.role == 'admin':
            user.role = 'user'
            db.session.commit()
            flash(f'Kullanıcı "{user.username}" normal kullanıcı olarak düşürüldü')
    else:
        flash('Kullanıcı bulunamadı')
    
    return redirect(url_for('profile'))

@app.route('/bulk_download', methods=['POST'])
def bulk_download():
    if not is_logged_in():
        return redirect(url_for('login'))
    
    doc_ids = request.form.getlist('doc_ids')
    if not doc_ids:
        flash('İndirmek için döküman seçilmedi')
        return redirect(url_for('upload_file'))
    
    # If only one document, download directly
    if len(doc_ids) == 1:
        return redirect(url_for('download_file', doc_id=doc_ids[0]))
    
    # For multiple documents, create a ZIP file
    import zipfile
    import tempfile
    
    # Create a temporary ZIP file
    temp_zip = tempfile.NamedTemporaryFile(delete=False, suffix='.zip')
    
    try:
        with zipfile.ZipFile(temp_zip.name, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for doc_id in doc_ids:
                doc = Document.query.get(int(doc_id))
                if doc and doc.file_data:
                    # Add file to ZIP with its original name
                    zipf.writestr(doc.title, doc.file_data)
        
        # Send the ZIP file
        return send_file(
            temp_zip.name,
            download_name=f'dokumanlar_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.zip',
            as_attachment=True,
            mimetype='application/zip'
        )
    except Exception as e:
        flash(f'İndirme oluşturma hatası: {str(e)}')
        return redirect(url_for('upload_file'))
    finally:
        # Clean up temp file after sending
        try:
            os.unlink(temp_zip.name)
        except:
            pass

@app.route('/bulk_delete', methods=['POST'])
def bulk_delete():
    if not is_admin():
        flash('Yetkisiz erişim')
        return redirect(url_for('upload_file'))
    
    doc_ids = request.form.getlist('doc_ids')
    if not doc_ids:
        flash('Silmek için döküman seçilmedi')
        return redirect(url_for('upload_file'))
    
    try:
        deleted_count = 0
        for doc_id in doc_ids:
            doc = Document.query.get(int(doc_id))
            if doc:
                db.session.delete(doc)
                deleted_count += 1
        
        db.session.commit()
        flash(f'{deleted_count} döküman başarıyla silindi')
    except Exception as e:
        db.session.rollback()
        flash(f'Döküman silme hatası: {str(e)}')
    
    return redirect(url_for('upload_file'))

def send_email(to_email, subject, body):
    """Send email using SMTP"""
    try:
        # Debug için konsola da yazdır
        print(f"=== E-POSTA GÖNDERİLİYOR ===")
        print(f"Kime: {to_email}")
        print(f"Konu: {subject}")
        print(f"========================")
        
        # Gerçek e-posta gönderimi
        msg = MIMEMultipart()
        msg['From'] = app.config['MAIL_DEFAULT_SENDER']
        msg['To'] = to_email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'html'))
        
        server = smtplib.SMTP(app.config['MAIL_SERVER'], app.config['MAIL_PORT'])
        server.starttls()
        server.login(app.config['MAIL_USERNAME'], app.config['MAIL_PASSWORD'])
        text = msg.as_string()
        server.sendmail(app.config['MAIL_DEFAULT_SENDER'], to_email, text)
        server.quit()
        
        print("✅ E-posta başarıyla gönderildi!")
        return True
    except Exception as e:
        print(f"❌ E-posta gönderme hatası: {e}")
        return False

def generate_reset_token():
    """Generate a secure reset token"""
    return secrets.token_urlsafe(32)

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email'].strip()
        user = User.query.filter_by(email=email).first()
        
        if user:
            # Generate reset token
            reset_token = generate_reset_token()
            user.reset_token = reset_token
            user.reset_token_expires = datetime.datetime.now() + datetime.timedelta(hours=1)  # 1 saat geçerli
            
            try:
                db.session.commit()
                
                # Send reset email
                reset_url = url_for('reset_password', token=reset_token, _external=True)
                
                email_body = f"""
                <html>
                <body>
                    <h2>DMS - Şifre Sıfırlama</h2>
                    <p>Merhaba {user.username},</p>
                    <p>Şifrenizi sıfırlamak için aşağıdaki linke tıklayın:</p>
                    <p><a href="{reset_url}" style="background-color: #007bff; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">Şifremi Sıfırla</a></p>
                    <p>Bu link 1 saat boyunca geçerlidir.</p>
                    <p>Eğer şifre sıfırlama talebinde bulunmadıysanız, bu e-postayı göz ardı edebilirsiniz.</p>
                    <br>
                    <p>DMS Ekibi</p>
                </body>
                </html>
                """
                
                if send_email(user.email, "DMS - Şifre Sıfırlama", email_body):
                    flash('Şifre sıfırlama linki e-posta adresinize gönderildi.')
                else:
                    flash('E-posta gönderilirken bir hata oluştu. Lütfen tekrar deneyin.')
                    
            except Exception as e:
                db.session.rollback()
                flash('Bir hata oluştu. Lütfen tekrar deneyin.')
        else:
            # Güvenlik için, e-posta bulunamasa bile başarılı mesaj göster
            flash('Eğer bu e-posta adresi sistemde kayıtlı ise, şifre sıfırlama linki gönderildi.')
        
        return redirect(url_for('login'))
    
    return render_template('forgot_password.html')

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    user = User.query.filter_by(reset_token=token).first()
    
    if not user or not user.reset_token_expires or user.reset_token_expires < datetime.datetime.now():
        flash('Geçersiz veya süresi dolmuş sıfırlama linki.')
        return redirect(url_for('forgot_password'))
    
    if request.method == 'POST':
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']
        
        if new_password != confirm_password:
            flash('Şifreler eşleşmiyor!')
            return render_template('reset_password.html', token=token)
        
        if len(new_password) < 6:
            flash('Şifre en az 6 karakter olmalıdır!')
            return render_template('reset_password.html', token=token)
        
        # Update password and clear reset token
        user.password_hash = generate_password_hash(new_password)
        user.reset_token = None
        user.reset_token_expires = None
        
        try:
            db.session.commit()
            flash('Şifreniz başarıyla güncellendi! Şimdi giriş yapabilirsiniz.')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            flash('Şifre güncellenirken bir hata oluştu.')
    
    return render_template('reset_password.html', token=token)

if __name__ == '__main__':
    # Create database tables if they don't exist
    with app.app_context():
        db.create_all()
        print("✅ Database tables created/checked")
    
    app.run(debug=True)
